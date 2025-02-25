import streamlit as st  
import pandas as pd  
import numpy as np  
import plotly.express as px  
from scipy.special import factorial  
from sklearn.linear_model import LogisticRegression  
from sklearn.ensemble import RandomForestClassifier  
from xgboost import XGBClassifier  
from io import BytesIO  
from docx import Document  

# Fonction pour prédire avec le modèle Poisson  
def poisson_prediction(goals):  
    probabilities = []  
    for k in range(6):  # Calculer pour 0 à 5 buts  
        prob = np.exp(-goals) * (goals ** k) / factorial(k)  
        probabilities.append(prob)  
    return probabilities  

# Fonction pour créer un document Word avec les résultats  
def create_doc(results):  
    doc = Document()  
    doc.add_heading('Analyse de Matchs de Football et Prédictions de Paris Sportifs', level=1)  

    # Ajout des données des équipes  
    doc.add_heading('Données des Équipes', level=2)  
    doc.add_paragraph(f"Équipe Domicile: {results['Équipe Domicile']}")  
    doc.add_paragraph(f"Équipe Extérieure: {results['Équipe Extérieure']}")  
    doc.add_paragraph(f"Buts Prédit Domicile: {results['Buts Prédit Domicile']:.2f}")  
    doc.add_paragraph(f"Buts Prédit Extérieur: {results['Buts Prédit Extérieur']:.2f}")  

    # Ajout des probabilités des modèles  
    doc.add_heading('Probabilités des Modèles', level=2)  
    doc.add_paragraph(f"Probabilité Domicile: {results['Probabilité Domicile']:.2f}")  
    doc.add_paragraph(f"Probabilité Nul: {results['Probabilité Nul']:.2f}")  
    doc.add_paragraph(f"Probabilité Extérieure: {results['Probabilité Extérieure']:.2f}")  

    # Enregistrement du document  
    buffer = BytesIO()  
    doc.save(buffer)  
    buffer.seek(0)  
    return buffer  

# Fonction pour entraîner et prédire avec les modèles  
@st.cache_resource  
def train_models():  
    # Créer un ensemble de données d'entraînement synthétique une seule fois  
    np.random.seed(42)  # Pour la reproductibilité  
    data = pd.DataFrame({  
        'home_goals': np.random.randint(0, 3, size=1000),  
        'away_goals': np.random.randint(0, 3, size=1000),  
        'xG': np.random.uniform(0, 2, size=1000),  
        'xGA': np.random.uniform(0, 2, size=1000),  
        'encais': np.random.uniform(0, 2, size=1000),  
        'result': np.random.choice([0, 1, 2], size=1000)  
    })  

    # Séparer les caractéristiques et la cible  
    X = data[['home_goals', 'away_goals', 'xG', 'xGA', 'encais']]  
    y = data['result']  

    # Modèle de régression logistique  
    log_reg = LogisticRegression(max_iter=50)  # Réduire le nombre d'itérations pour la démonstration  
    log_reg.fit(X, y)  

    # Modèle Random Forest  
    rf = RandomForestClassifier(n_estimators=50)  # Réduire le nombre d'estimations pour la démonstration  
    rf.fit(X, y)  

    # Modèle XGBoost  
    xgb = XGBClassifier(use_label_encoder=False, eval_metric='logloss', n_estimators=50)  # Réduire le nombre d'estimations  
    xgb.fit(X, y)  

    return log_reg, rf, xgb  

# Vérifier si les modèles sont déjà chargés dans l'état de session  
if 'models' not in st.session_state:  
    st.session_state.models = train_models()  

log_reg_model, rf_model, xgb_model = st.session_state.models  

# Interface utilisateur  
st.title("🏆 Analyse de Matchs de Football et Prédictions de Paris Sportifs")  

# Saisie des données des équipes  
st.header("Saisie des données des équipes")  

# Utilisation d'accordéons pour les statistiques des équipes  
with st.expander("Statistiques de l'équipe à domicile", expanded=True):  
    if 'home_team' not in st.session_state:  
        st.session_state.home_team = "Équipe A"  
    st.session_state.home_team = st.text_input("Nom de l'équipe à domicile", value=st.session_state.home_team)  

    # Variables pour l'équipe à domicile  
    home_vars = [  
        "goals", "xG", "encais", "xGA", "tirs_par_match", "passes_menant_a_tir",  
        "tirs_cadres", "tirs_concedes", "duels_defensifs", "possession",  
        "passes_reussies", "touches_surface", "forme_recente"  
    ]  
    for var in home_vars:  
        if f'home_{var}' not in st.session_state:  
            st.session_state[f'home_{var}'] = 0.0  
        st.session_state[f'home_{var}'] = st.slider(  
            f"{var.replace('_', ' ').title()} (domicile)",  
            min_value=0.0, max_value=5.0 if var in ["goals", "xG", "encais", "xGA"] else 100.0,  
            value=st.session_state[f'home_{var}'], step=0.1  
        )  

with st.expander("Statistiques de l'équipe à l'extérieur", expanded=True):  
    if 'away_team' not in st.session_state:  
        st.session_state.away_team = "Équipe B"  
    st.session_state.away_team = st.text_input("Nom de l'équipe à l'extérieur", value=st.session_state.away_team)  

    # Variables pour l'équipe à l'extérieur  
    for var in home_vars:  
        if f'away_{var}' not in st.session_state:  
            st.session_state[f'away_{var}'] = 0.0  
        st.session_state[f'away_{var}'] = st.slider(  
            f"{var.replace('_', ' ').title()} (extérieur)",  
            min_value=0.0, max_value=5.0 if var in ["goals", "xG", "encais", "xGA"] else 100.0,  
            value=st.session_state[f'away_{var}'], step=0.1  
        )  

# Saisie des cotes des bookmakers  
st.header("Cotes des Équipes")  
if 'odds_home' not in st.session_state:  
    st.session_state.odds_home = 1.8  
if 'odds_away' not in st.session_state:  
    st.session_state.odds_away = 2.2  
st.session_state.odds_home = st.number_input("Cote pour l'équipe à domicile", min_value=1.0, value=st.session_state.odds_home)  
st.session_state.odds_away = st.number_input("Cote pour l'équipe à l'extérieur", min_value=1.0, value=st.session_state.odds_away)  

# Calcul des probabilités implicites  
def calculate_implied_prob(odds):  
    return 1 / odds  

# Prédictions  
if st.button("🔍 Prédire les résultats"):  
    # Calcul des buts prédit  
    home_goals_pred = st.session_state.home_goals + st.session_state.home_xG - st.session_state.away_encais  
    away_goals_pred = st.session_state.away_goals + st.session_state.away_xG - st.session_state.home_encais  

    # Calcul des probabilités avec le modèle de Poisson  
    home_probabilities = poisson_prediction(home_goals_pred)  
    away_probabilities = poisson_prediction(away_goals_pred)  

    # Formatage des résultats pour l'affichage  
    home_results = ", ".join([f"{i} but {home_probabilities[i] * 100:.1f}%" for i in range(len(home_probabilities))])  
    away_results = ", ".join([f"{i} but {away_probabilities[i] * 100:.1f}%" for i in range(len(away_probabilities))])  

    # Calcul des probabilités implicites  
    implied_home_prob = calculate_implied_prob(st.session_state.odds_home)  
    implied_away_prob = calculate_implied_prob(st.session_state.odds_away)  
    implied_draw_prob = 1 - (implied_home_prob + implied_away_prob)  

    # Prédictions avec les modèles  
    input_data = [[home_goals_pred, away_goals_pred, st.session_state.home_xG, st.session_state.away_xG, st.session_state.home_encais, st.session_state.away_encais]]  

    try:  
        log_reg_prob = log_reg_model.predict_proba(input_data)[0]  
        rf_prob = rf_model.predict_proba(input_data)[0]  
        xgb_prob = xgb_model.predict_proba(input_data)[0]  
    except Exception as e:  
        st.error(f"Erreur lors de la prédiction : {e}")  
        log_reg_prob, rf_prob, xgb_prob = None, None, None  

    # Affichage des résultats  
    st.subheader("📊 Résultats des Prédictions")  

    # Tableau pour le modèle de Poisson  
    poisson_results = pd.DataFrame({  
        "Équipe": [st.session_state.home_team, st.session_state.away_team],  
        "Buts Prédit": [home_results, away_results]  
    })  

    st.markdown("### Résultats du Modèle de Poisson")  
    st.dataframe(poisson_results, use_container_width=True)  

    # Détails sur chaque prédiction des modèles  
    st.markdown("### Détails des Prédictions des Modèles")  
    model_details = {  
        "Modèle": ["Régression Logistique", "Random Forest", "XGBoost"],  
        "Victoire Domicile (%)": [log_reg_prob[2] * 100 if log_reg_prob is not None else None,  
                                  rf_prob[2] * 100 if rf_prob is not None else None,  
                                  xgb_prob[2] * 100 if xgb_prob is not None else None],  
        "Match Nul (%)": [log_reg_prob[1] * 100 if log_reg_prob is not None else None,  
                          rf_prob[1] * 100 if rf_prob is not None else None,  
                          xgb_prob[1] * 100 if xgb_prob is not None else None],  
        "Victoire Extérieure (%)": [log_reg_prob[0] * 100 if log_reg_prob is not None else None,  
                                    rf_prob[0] * 100 if rf_prob is not None else None,  
                                    xgb_prob[0] * 100 if xgb_prob is not None else None],  
    }  
    model_details_df = pd.DataFrame(model_details)  
    st.dataframe(model_details_df, use_container_width=True)  

    # Comparaison des probabilités implicites et prédites  
    st.subheader("📊 Comparaison des Probabilités Implicites et Prédites")  
    comparison_data = {  
        "Type": ["Implicite Domicile", "Implicite Nul", "Implicite Extérieure", "Prédite Domicile", "Prédite Nul", "Prédite Extérieure"],  
        "Probabilité (%)": [  
            implied_home_prob * 100,  
            implied_draw_prob * 100,  
            implied_away_prob * 100,  
            log_reg_prob[2] * 100 if log_reg_prob is not None else None,  
            log_reg_prob[1] * 100 if log_reg_prob is not None else None,  
            log_reg_prob[0] * 100 if log_reg_prob is not None else None,  
        ]  
    }  
    comparison_df = pd.DataFrame(comparison_data)  
    st.dataframe(comparison_df, use_container_width=True)  

    # Graphique des performances des modèles  
    st.subheader("📈 Comparaison des Modèles")  
    model_comparison_data = {  
        "Modèle": ["Régression Logistique", "Random Forest", "XGBoost"],  
        "Probabilité Domicile (%)": [log_reg_prob[2] * 100 if log_reg_prob is not None else 0,  
                                      rf_prob[2] * 100 if rf_prob is not None else 0,  
                                      xgb_prob[2] * 100 if xgb_prob is not None else 0],  
        "Probabilité Nul (%)": [log_reg_prob[1] * 100 if log_reg_prob is not None else 0,  
                                rf_prob[1] * 100 if rf_prob is not None else 0,  
                                xgb_prob[1] * 100 if xgb_prob is not None else 0],  
        "Probabilité Extérieure (%)": [log_reg_prob[0] * 100 if log_reg_prob is not None else 0,  
                                       rf_prob[0] * 100 if rf_prob is not None else 0,  
                                       xgb_prob[0] * 100 if xgb_prob is not None else 0],  
    }  
    model_comparison_df = pd.DataFrame(model_comparison_data)  
    fig = px.bar(model_comparison_df, x='Modèle', y=['Probabilité Domicile (%)', 'Probabilité Nul (%)', 'Probabilité Extérieure (%)'],  
                  title='Comparaison des Probabilités des Modèles', barmode='group')  
    st.plotly_chart(fig)  

    # Explication des modèles  
    st.subheader("📊 Explication des Modèles")  
    st.write("""  
    - **Régression Logistique** : Modèle utilisé pour prédire la probabilité d'un événement binaire.  
    - **Random Forest** : Modèle d'ensemble qui utilise plusieurs arbres de décision pour améliorer la précision.  
    - **XGBoost** : Modèle d'apprentissage par boosting qui est très efficace pour les compétitions de machine learning.  
    """)  

    # Option de téléchargement des résultats  
    results = {  
        "Équipe Domicile": st.session_state.home_team,  
        "Équipe Extérieure": st.session_state.away_team,  
        "Buts Prédit Domicile": home_goals_pred,  
        "Buts Prédit Extérieur": away_goals_pred,  
        "Probabilité Domicile": log_reg_prob[2] if log_reg_prob is not None else None,  
        "Probabilité Nul": log_reg_prob[1] if log_reg_prob is not None else None,  
        "Probabilité Extérieure": log_reg_prob[0] if log_reg_prob is not None else None,  
    }  

    if st.button("📥 Télécharger les résultats en DOC"):  
        buffer = create_doc(results)  
        st.download_button(  
            label="Télécharger les résultats",  
            data=buffer,  
            file_name="predictions.docx",  
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"  
        )
